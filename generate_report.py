import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from analysis import load_data, get_kpis, get_sales_by_division, get_regional_performance, get_top_products

def generate_report():
    df = load_data()
    kpis = get_kpis(df)
    div_df = get_sales_by_division(df)
    reg_df = get_regional_performance(df)
    top_p = get_top_products(df, 5)

    doc = docx.Document()

    # Title
    title = doc.add_heading('Nassau Candy Distributor - Executive Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Comprehensive financial performance, regional breakdown, and product analytics report.')
    doc.add_paragraph()

    # Section 1: Executive Summary & KPIs
    doc.add_heading('1. Executive Summary & Core KPIs', level=1)
    p = doc.add_paragraph()
    p.add_run(f"• Total Revenue Generated: ").bold = True
    p.add_run(f"${kpis['total_sales']:,.2f}\n")
    p.add_run(f"• Total Gross Profit: ").bold = True
    p.add_run(f"${kpis['total_profit']:,.2f}\n")
    p.add_run(f"• Overall Profit Margin: ").bold = True
    p.add_run(f"{kpis['overall_margin']:.2f}%\n")
    p.add_run(f"• Total Units Sold: ").bold = True
    p.add_run(f"{kpis['total_units']:,}\n")
    p.add_run(f"• Total Unique Orders: ").bold = True
    p.add_run(f"{kpis['total_orders']:,}")

    # Section 2: Division Performance
    doc.add_heading('2. Performance by Division', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Division'
    hdr_cells[1].text = 'Sales ($)'
    hdr_cells[2].text = 'Gross Profit ($)'
    hdr_cells[3].text = 'Units Sold'

    for index, row in div_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Division'])
        row_cells[1].text = f"${row['Sales']:,.2f}"
        row_cells[2].text = f"${row['Gross Profit']:,.2f}"
        row_cells[3].text = f"{int(row['Units']):,}"

    # Section 3: Regional Performance
    doc.add_heading('3. Regional Breakdown', level=1)
    reg_table = doc.add_table(rows=1, cols=4)
    reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_hdr_cells = reg_table.rows[0].cells
    r_hdr_cells[0].text = 'Region'
    r_hdr_cells[1].text = 'Sales ($)'
    r_hdr_cells[2].text = 'Gross Profit ($)'
    r_hdr_cells[3].text = 'Units Sold'

    for index, row in reg_df.iterrows():
        r_row_cells = reg_table.add_row().cells
        r_row_cells[0].text = str(row['Region'])
        r_row_cells[1].text = f"${row['Sales']:,.2f}"
        r_row_cells[2].text = f"${row['Gross Profit']:,.2f}"
        r_row_cells[3].text = f"{int(row['Units']):,}"

    # Section 4: Top 5 Products
    doc.add_heading('4. Top 5 Products by Revenue', level=1)
    top_table = doc.add_table(rows=1, cols=3)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_hdr_cells = top_table.rows[0].cells
    t_hdr_cells[0].text = 'Product Name'
    t_hdr_cells[1].text = 'Sales ($)'
    t_hdr_cells[2].text = 'Gross Profit ($)'

    for index, row in top_p.iterrows():
        t_row_cells = top_table.add_row().cells
        t_row_cells[0].text = str(row['Product Name'])
        t_row_cells[1].text = f"${row['Sales']:,.2f}"
        t_row_cells[2].text = f"${row['Gross Profit']:,.2f}"

    doc.save('report.docx')
    print("report.docx generated successfully!")

if __name__ == '__main__':
    generate_report()
